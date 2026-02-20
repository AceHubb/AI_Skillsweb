import json
import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

def get_cards_without_pdfs(cards_file, pdf_mapping_file, limit=10):
    with open(cards_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    with open(pdf_mapping_file, 'r', encoding='utf-8') as f:
        pdf_mapping = json.load(f)
    
    mapped_ids = set(pdf_mapping.values())
    
    target_cards = []
    for card in data['cards']:
        card_id = card['id']
        
        # Check if PDF exists in 'media' or 'pdf_mapping'
        has_pdf = False
        if card_id in mapped_ids:
            has_pdf = True
        
        if 'media' in card:
            for item in card['media']:
                if item.endswith('.pdf'):
                    has_pdf = True
                    break
        
        if not has_pdf:
            target_cards.append(card)
            if len(target_cards) >= limit:
                break
                
    return target_cards

def generate_strategic_content(card, section):
    title = card.get('title', 'Unknown Asset')
    desc = card.get('description', 'No description available.')
    
    # Template content based on UK English Production Order
    # Note: Using UK English spellings: Modernisation, Programme, Initialisation, Honour, Organisation, etc.
    
    templates = {
        "Executive Summary": (
            f"The '{title}' asset represents a critical component within the overarching 30-year engineering lifecycle. "
            f"Its strategic integration ensures the modernisation of our technical infrastructure, "
            f"aligning with the long-term programme for enterprise excellence. {desc}"
        ),
        "Operational Continuity": (
            "Maintaining the highest standards of reliability and stability is paramount for this asset. "
            "System health monitoring is integrated into the operational framework to mitigate risk and ensure "
            "seamless continuity of service across the global infrastructure."
        ),
        "Governance & Compliance": (
            "Strict adherence to regulatory requirements and internal security protocols is mandatory. "
            "Risk management processes are prioritised to ensure full alignment with UK and international standards, "
            "upholding the integrity of the data and the security of the broader ecosystem."
        ),
        "Strategic ROI": (
            "The initialisation of this asset delivers measurable business value through enhanced cost-efficiency. "
            "Strategic alignment with organisational growth objectives ensures that the investment contributes "
            "directly to the company's long-term financial health and competitive advantage."
        ),
        "Integration Roadmap": (
            "Future-proofing is at the heart of the integration roadmap. The asset is designed for scalability, "
            "ensuring that as technology evolves, our systems remain robust and capable of supporting "
            "emerging requirements without the need for total re-engineering."
        )
    }
    return templates.get(section, "")

def create_dossier(target_cards, output_file="EXECUTIVE_DOSSIER.docx"):
    doc = Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for i, card in enumerate(target_cards):
        if i > 0:
            doc.add_page_break()
        
        # Hidden/Footer-like text at start of page (as per prompt order, though usually footer is better)
        # We will put it at the very top as small light gray text as "hidden" metadata
        p_metadata = doc.add_paragraph()
        run_m = p_metadata.add_run(f"NEW CARD: [{card['id']}]")
        run_m.font.size = Pt(8)
        run_m.font.color.rgb = None # or light gray
        
        # Main Header
        doc.add_heading(card['title'], level=1)
        
        sections = [
            "Executive Summary",
            "Operational Continuity",
            "Governance & Compliance",
            "Strategic ROI",
            "Integration Roadmap"
        ]
        
        for sec in sections:
            doc.add_heading(sec, level=2)
            content = generate_strategic_content(card, sec)
            doc.add_paragraph(content)
            
        # Add some spacing to encourage 2-page length if needed, 
        # but for now we follow the structure.
        # Professional padding
        for _ in range(3):
            doc.add_paragraph("")

    doc.save(output_file)
    print(f"Dossier generated: {output_file}")

if __name__ == "__main__":
    cards_path = "cards.json"
    mapping_path = "pdf_mapping.json"
    
    if not os.path.exists(cards_path) or not os.path.exists(mapping_path):
        print("Required files missing.")
        sys.exit(1)
        
    targets = get_cards_without_pdfs(cards_path, mapping_path, limit=10)
    if not targets:
        print("No cards without PDFs found.")
    else:
        create_dossier(targets)
