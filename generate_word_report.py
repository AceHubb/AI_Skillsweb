import sys
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report(text_file, output_file=None, logo_path=None):
    if not os.path.exists(text_file):
        print(f"Error: File {text_file} not found.")
        return

    if output_file is None:
        output_file = text_file.replace('.txt', '.docx')
        if output_file == text_file:
            output_file = text_file + '.docx'

    doc = Document()

    # Add Logo if provided
    if logo_path and os.path.exists(logo_path):
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(2.0))
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    with open(text_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Simple parser for the report structure
    for i, line in enumerate(lines):
        line = line.strip('\n')
        
        # Header 1 (Title)
        if i == 0 and lines[i+1].startswith('===='):
            heading = doc.add_heading(line, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        if line.startswith('===='):
            continue
            
        # Header 2
        if i > 0 and i < len(lines)-1 and lines[i+1].startswith('----'):
            doc.add_heading(line, level=1)
            continue
            
        if line.startswith('----'):
            continue

        # Tree Diagram or List Items
        if line.lstrip().startswith(('└──', '├──', '│', '└──', '* ')):
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            # Use courier for tree diagrams
            if any(char in line for char in ['└', '├', '│', '─']):
                run = p.runs[0] if p.runs else p.add_run(line)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue

        # Regular text
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    doc.save(output_file)
    print(f"Report generated: {output_file}")
    return output_file

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_word_report.py <text_file> [logo_path]")
    else:
        file_to_convert = sys.argv[1]
        branding_logo = sys.argv[2] if len(sys.argv) > 2 else "Waalbridge.png"
        create_word_report(file_to_convert, logo_path=branding_logo)
